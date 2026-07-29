"""Build the submission-ready Word document from the annotated thesis HTML.

Pandoc converts the HTML directly but produces a document with pandoc's styles
rather than a thesis template's: no live table of contents, no page numbering,
tables without a Word table style, references without a hanging indent, and the
two-column annotation layout flattened into an unreadable interleaving. Since
the document is opened, edited and submitted in Word, it is built here with
python-docx so every element carries a real Word style.

The source is an ANNOTATED thesis: each section holds a thesis panel and an
annotation panel explaining the writing choices. Word has no side-by-side
equivalent, so annotations are emitted after each section's thesis text as a
clearly separated commentary table. Pass --no-annotations for a submission-only
document containing the thesis text alone.

Charts are inline SVG and do not survive any HTML to DOCX path, so each is
rebuilt as a native Word table carrying the same figures under the same caption.

Run:  python scripts/build_thesis_docx.py thesis_source.html
      python scripts/build_thesis_docx.py thesis_source.html --no-annotations
"""
from __future__ import annotations

import argparse
import html
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PNG_DIR = Path('reports/figures/thesis_png')

BODY_FONT = "Times New Roman"
SANS = "Calibri"
ACCENT = RGBColor(0x9A, 0x34, 0x12)
MUTED = RGBColor(0x59, 0x55, 0x4F)

# Highlight spans are reading aids for the web page, not emphasis in the thesis
# text. Carrying them into Word as bold made 25% of all runs bold, which reads as
# shouting in a printed document, so they are unwrapped to plain text instead.
HG = re.compile(r'<span class="hg hg-[a-z]+">(.*?)</span>', re.S)


# --------------------------------------------------------------- word plumbing

def _field(paragraph, instruction: str, placeholder: str):
    run = paragraph.add_run()
    for tag, attr, text in (("w:fldChar", "begin", None), ("w:instrText", None, instruction),
                            ("w:fldChar", "separate", None), ("w:t", None, placeholder),
                            ("w:fldChar", "end", None)):
        el = OxmlElement(tag)
        if attr:
            el.set(qn("w:fldCharType"), attr)
        if text is not None:
            el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)


def _configure(doc: Document):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.right_margin = Inches(1)
        s.left_margin = Inches(1.25)

    n = doc.styles["Normal"]
    n.font.name = BODY_FONT
    n.font.size = Pt(12)
    n.paragraph_format.line_spacing = 1.5
    n.paragraph_format.space_after = Pt(8)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    for name, size, colour in [("Heading 1", 18, ACCENT), ("Heading 2", 14, None),
                               ("Heading 3", 12.5, None)]:
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = colour or RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(16 if size == 18 else 12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.15
        st.paragraph_format.keep_with_next = True

    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(p, "PAGE", "1")
    for r in p.runs:
        r.font.name = BODY_FONT
        r.font.size = Pt(10)


# ------------------------------------------------------------------- html read

def _text(s: str) -> str:
    s = re.sub(r"<br\s*/?>", " ", s)
    s = re.sub(r"<[^>]+>", "", s)
    return re.sub(r"\s+", " ", html.unescape(s)).strip()


def _runs(par, frag: str, base_size: float | None = None, sans: bool = False):
    frag = HG.sub(r"\1", frag)
    frag = re.sub(r"<br\s*/?>", " ", frag)
    parts = re.split(r"(<(?:strong|b|em|i|code)[^>]*>.*?</(?:strong|b|em|i|code)>)", frag, flags=re.S)
    for part in parts:
        if not part:
            continue
        m = re.match(r"<(strong|b|em|i|code)[^>]*>(.*?)</\1>", part, flags=re.S)
        if m:
            tag, inner = m.group(1), _text(m.group(2))
            if not inner:
                continue
            r = par.add_run(inner)
            r.bold = tag in ("strong", "b")
            r.italic = tag in ("em", "i")
            if tag == "code":
                r.font.name = "Consolas"
                r.font.size = Pt(10)
            else:
                if base_size:
                    r.font.size = Pt(base_size)
                if sans:
                    r.font.name = SANS
        else:
            # Preserve the single space that sits BETWEEN inline elements. A full
            # strip here silently welds runs together, producing "Question 2.What"
            # wherever a bold label is followed by italic text.
            raw = html.unescape(re.sub(r"<[^>]+>", "", re.sub(r"<br\s*/?>", " ", part)))
            txt = re.sub(r"\s+", " ", raw)
            if not txt.strip():
                if txt and par.runs:
                    par.add_run(" ")
                continue
            r = par.add_run(txt)
            if base_size:
                r.font.size = Pt(base_size)
            if sans:
                r.font.name = SANS


# ------------------------------------------------------------------ front page

def _front(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    rows = [("CONTESTED SHOT QUALITY", True, 26, False),
            ("A calibrated, leakage-audited machine learning system for context-aware "
             "basketball shot prediction, and a measured account of where its ceiling lies",
             False, 13, True),
            ("", False, 12, False), ("", False, 12, False),
            ("Nischay Maharjan", True, 14, False),
            ("Student Number 230357", False, 12, False), ("", False, 12, False),
            ("BSc (Hons) Computing with Artificial Intelligence", False, 12, False),
            ("Individual Research Project", False, 12, False), ("", False, 12, False),
            ("Softwarica College of IT and E-Commerce", False, 12, False),
            ("in academic partnership with Coventry University, United Kingdom", False, 12, False),
            ("", False, 12, False), ("2026", False, 12, False)]
    for text, bold, size, italic in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6 if size > 20 else 4)
        if text:
            r = p.add_run(text)
            r.bold, r.italic = bold, italic
            r.font.size = Pt(size)
            r.font.name = BODY_FONT
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_paragraph("Table of Contents", style="Heading 1")
    p = doc.add_paragraph()
    _field(p, r'TOC \o "1-3" \h \z \u', "Right-click and choose Update Field")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# --------------------------------------------------------------- infographics

def _chart(doc: Document, title: str, headers: list[str], rows: list[tuple]):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = SANS
    p.paragraph_format.space_after = Pt(3)

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = SANS
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            par = cells[i].paragraphs[0]
            par.paragraph_format.line_spacing = 1.0
            par.paragraph_format.space_after = Pt(2)
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = SANS


CHARTS = {
 "Infographic 1": ("Murphy partition of the Brier score", ["Component", "Value", "Meaning"],
    [("Uncertainty", "0.2492", "Irreducible, fixed by the 47.1% base rate"),
     ("Resolution", "0.0349", "What the model extracts from context"),
     ("Reliability", "0.0001", "Miscalibration, essentially zero"),
     ("Reconstructed", "0.2144", "Measured Brier 0.2137, residual 0.0007"),
     ("Skill, model", "0.1425", "Brier Skill Score against the base rate"),
     ("Skill, baseline", "0.0682", "Slightly less than half the model's")]),
 "Infographic 2": ("Model families on the validation season", ["Model", "AUC", "Brier", "Train (s)"],
    [("CatBoost, selected", "0.6976", "0.2133", "117"), ("XGBoost", "0.6970", "0.2134", "35"),
     ("FT-Transformer", "0.6967", "0.2138", "3,054"), ("LightGBM", "0.6963", "0.2135", "82"),
     ("Tabular MLP", "0.6920", "0.2151", "455"), ("Random forest", "0.6894", "0.2173", "345"),
     ("TabNet", "0.6856", "0.2170", "365"), ("Logistic regression", "0.6707", "0.2229", "19"),
     ("xP baseline", "0.6335", "n/a", "n/a")]),
 "Infographic 3": ("Predictability of post-release flight", ["Quantity", "R squared", "Interpretation"],
    [("Apex height", "0.612", "Predictable, ballistic envelope"),
     ("Flight time", "0.561", "Predictable, ballistic envelope"),
     ("Minimum distance to rim", "0.363", "Weakly predictable, decides the outcome"),
     ("Entry angle", "0.046", "Essentially unpredictable, decides the outcome")]),
 "Infographic 4": ("The same pipeline on two targets", ["Target", "Unit", "n", "AUC"],
    [("Single shot", "One attempt", "219,157", "0.7001"),
     ("Player season", "Player-season", "1,172", "0.8099")]),
 "Infographic 5": ("Discrimination by court zone", ["Zone", "AUC", "n"],
    [("Restricted area", "0.7247", "62,250"), ("Paint, non-RA", "0.6349", "43,908"),
     ("Mid-range", "0.6130", "22,025"), ("Right corner three", "0.6047", "11,360"),
     ("Left corner three", "0.6031", "12,210"), ("Above the break three", "0.5953", "67,375"),
     ("Aggregate", "0.7001", "219,157")]),
 "Figure 1": ("The iterative research cycle", ["Stage", "Action"],
    [("1. Question", "One testable claim"), ("2. Implement", "Smallest change"),
     ("3. Measure", "Against the fixed validation season"),
     ("4. Decide", "Keep, or record the null"),
     ("Loop", "Every outcome logged; the test season is never touched inside it")]),
 "Figure 2": ("Layered system architecture", ["Layer", "Contents", "Discipline enforced"],
    [("1 Acquisition", "Shot records, play-by-play, API, tracking, profiles", "Validated against overlap"),
     ("2 Features", "Spatial, temporal, shot type, player, possession, skill", "Train-only statistics"),
     ("3 Modelling", "Eight families, isotonic calibration, frozen bundle", "Config hash recorded"),
     ("4 Serving", "Inference, explanation, ranking, contest, movement", "One prediction path"),
     ("5 Presentation", "Dashboard, 3D scenario builder", "No modelling logic")]),
}


_CHART_CACHE: dict[str, Path] = {}


def _figure(doc: Document, caption: str):
    """Insert the rendered figure.

    Keyed on an exact "Figure N" match rather than a prefix, because a prefix
    test lets "Figure 1" capture "Figure 10" through "Figure 19" and silently
    serve the wrong image under the right caption. Falls back to a data table
    where one is defined, since a missing figure is worse than an unlovely one.
    """
    m = re.match(r"\s*(Figure|Table)\s*(\d+)", caption)
    if m:
        png = PNG_DIR / f"{m.group(1)}{m.group(2)}.png"
        if png.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            p.add_run().add_picture(str(png), width=Inches(6.0))
            return
    for key, (title, heads, rows) in CHARTS.items():
        if caption.startswith(key):
            _chart(doc, title, heads, rows)
            return


def _caption(doc: Document, text: str):
    if not text:
        return
    p = doc.add_paragraph()
    m = re.match(r"^((?:Figure|Table|Infographic)\s*[\d]+)\.?\s*(.*)$", text, flags=re.S)
    if m:
        a = p.add_run(m.group(1) + "  ")
        a.bold = True
        b = p.add_run(m.group(2))
        b.italic = True
        for r in (a, b):
            r.font.size = Pt(9.5)
            r.font.name = SANS
    else:
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(12)


def _table(doc: Document, chunk: str):
    heads = re.findall(r"<th[^>]*>(.*?)</th>", chunk, flags=re.S)
    body = chunk[chunk.index("<tbody>"):] if "<tbody>" in chunk else chunk
    rows = [re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", r, flags=re.S)
            for r in re.findall(r"<tr[^>]*>(.*?)</tr>", body, flags=re.S)]
    rows = [r for r in rows if r]
    if not rows:
        return
    ncols = max(len(heads), max(len(r) for r in rows))
    t = doc.add_table(rows=1 if heads else 0, cols=ncols)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if heads:
        for i in range(ncols):
            r = t.rows[0].cells[i].paragraphs[0].add_run(_text(heads[i]) if i < len(heads) else "")
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = SANS
    for cells in rows:
        row = t.add_row()
        for i in range(ncols):
            par = row.cells[i].paragraphs[0]
            par.paragraph_format.line_spacing = 1.0
            par.paragraph_format.space_after = Pt(2)
            if i < len(cells):
                _runs(par, cells[i], base_size=9.5, sans=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ----------------------------------------------------------------- conversion

def _emit_panel(doc: Document, panel: str, is_refs: bool):
    if is_refs:
        for entry in re.findall(r"<p>(.*?)</p>", panel, flags=re.S):
            if not _text(entry):
                continue
            par = doc.add_paragraph()
            _runs(par, entry)
            par.paragraph_format.left_indent = Inches(0.5)
            par.paragraph_format.first_line_indent = Inches(-0.5)
            par.paragraph_format.line_spacing = 1.0
            par.paragraph_format.space_after = Pt(9)
        return

    pos = 0
    for m in re.finditer(r'<(h4|p|ul|ol|pre|figure|div)\b[^>]*>.*?</\1>', panel, flags=re.S):
        if m.start() < pos:
            continue
        pos = m.end()
        tag, chunk = m.group(1), m.group(0)

        if tag == "h4":
            doc.add_paragraph(_text(chunk), style="Heading 2")
        elif tag == "p":
            inner = re.sub(r"^<p[^>]*>|</p>$", "", chunk, flags=re.S)
            if _text(inner):
                _runs(doc.add_paragraph(), inner)
        elif tag in ("ul", "ol"):
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in re.findall(r"<li[^>]*>(.*?)</li>", chunk, flags=re.S):
                _runs(doc.add_paragraph(style=style), li)
        elif tag == "pre":
            p = doc.add_paragraph()
            r = p.add_run(html.unescape(re.sub(r"<[^>]+>", "", chunk)).strip())
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.line_spacing = 1.0
        elif tag == "figure":
            cap = re.search(r"<figcaption>(.*?)</figcaption>", chunk, flags=re.S)
            cap_text = _text(cap.group(1)) if cap else ""
            _figure(doc, cap_text)
            _caption(doc, cap_text)
        elif tag == "div" and "tablewrap" in chunk:
            _table(doc, chunk)
        elif tag == "div" and "keyfig" in chunk:
            cells = re.findall(r'<span class="k">(.*?)</span><span class="v">(.*?)</span><span class="n">(.*?)</span>',
                               chunk, flags=re.S)
            if cells:
                _chart(doc, "Dataset split", ["Split", "Shots", "Seasons"],
                       [(_text(k), _text(v), _text(n)) for k, v, n in cells])


def _emit_annotations(doc: Document, ann: str):
    cards = re.findall(r'<div class="ac[^"]*"><div class="ac-type">(.*?)</div>\s*<p>(.*?)</p>', ann, flags=re.S)
    if not cards:
        return
    p = doc.add_paragraph()
    r = p.add_run("Annotation: writing and design decisions in this section")
    r.bold = r.italic = True
    r.font.size = Pt(10)
    r.font.name = SANS
    r.font.color.rgb = MUTED
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

    t = doc.add_table(rows=0, cols=2)
    t.style = "Light List Accent 1"
    for kind, text in cards:
        cells = t.add_row().cells
        a = cells[0].paragraphs[0].add_run(_text(kind))
        a.bold = True
        a.font.size = Pt(9)
        a.font.name = SANS
        b = cells[1].paragraphs[0]
        b.paragraph_format.line_spacing = 1.0
        _runs(b, text, base_size=9.5, sans=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build(source: Path, out: Path, annotations: bool) -> Path:
    raw = source.read_text(encoding="utf-8")
    doc = Document()
    _configure(doc)
    _front(doc)

    sections = re.findall(r'<section class="sec" id="([^"]+)">(.*?)</section>', raw, flags=re.S)
    if not sections:
        raise SystemExit("no sections found; is this the annotated source?")

    for sec_id, body in sections:
        tag = re.search(r'class="sec-tag">(.*?)</span>', body, flags=re.S)
        title = re.search(r'class="sec-title">(.*?)</h2>', body, flags=re.S)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        if tag:
            p = doc.add_paragraph()
            r = p.add_run(_text(tag.group(1)).upper())
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.name = SANS
            r.font.color.rgb = ACCENT
            p.paragraph_format.space_after = Pt(0)
        doc.add_paragraph(_text(title.group(1)) if title else sec_id, style="Heading 1")

        is_refs = 'thesis-panel refs' in body
        panel = re.search(r'<div class="thesis-panel[^"]*"[^>]*>(.*)', body, flags=re.S)
        if panel:
            content = panel.group(1)
            cut = content.find('<div class="ann-panel">')
            _emit_panel(doc, content[:cut] if cut > 0 else content, is_refs)
        elif "svgframe" in body:
            # section 01 carries an introduction and figures, but no thesis panel
            intro = re.search(r'<p style="[^"]*">(.*?)</p>', body, flags=re.S)
            if intro:
                _runs(doc.add_paragraph(), intro.group(1))
            for fig in re.findall(r"<figure>.*?</figure>", body, flags=re.S):
                cap = re.search(r"<figcaption>(.*?)</figcaption>", fig, flags=re.S)
                cap_text = _text(cap.group(1)) if cap else ""
                _figure(doc, cap_text)
                _caption(doc, cap_text)
            # this ran before the figures and was a literal no-op, so the
            # section opened straight into a chart with no introduction
            pass

        if annotations:
            ann = re.search(r'<div class="ann-panel">(.*)$', body, flags=re.S)
            if ann:
                _emit_annotations(doc, ann.group(1))

    doc.save(out)
    _assert_complete(raw, doc, annotations)
    return out


def _assert_complete(raw: str, doc: Document, annotations: bool):
    """Fail loudly if content in the HTML never reached the document.

    An earlier build silently dropped the entire reference list, because the
    block regex consumes a whole container and anything inside an unhandled one
    disappears without a trace. Nothing in the output signalled the loss, so the
    invariant is checked rather than assumed.
    """
    want_secs = len(re.findall(r'<section class="sec"', raw))
    got_secs = len([p for p in doc.paragraphs if p.style.name == "Heading 1"]) - 1
    refs_block = re.search(r'class="thesis-panel refs".*?(?=</section>)', raw, flags=re.S)
    want_refs = len(re.findall(r"<p>", refs_block.group(0))) if refs_block else 0
    got_refs = len([p for p in doc.paragraphs
                    if p.paragraph_format.first_line_indent
                    and p.paragraph_format.first_line_indent < 0])
    want_cards = len(re.findall(r'<div class="ac-type">', raw))
    got_cards = sum(len(t.rows) for t in doc.tables if t.style.name == "Light List Accent 1")

    # only <figure> elements become images; tables also carry a figcaption
    want_figs = len(re.findall(r"<figure>", raw))
    got_figs = len(doc.inline_shapes)

    bad = []
    if got_figs < want_figs:
        bad.append(f"figures {want_figs} -> {got_figs}")
    if got_secs < want_secs:
        bad.append(f"sections {want_secs} -> {got_secs}")
    if got_refs < want_refs:
        bad.append(f"references {want_refs} -> {got_refs}")
    if annotations and got_cards < want_cards:
        bad.append(f"annotations {want_cards} -> {got_cards}")
    if bad:
        raise SystemExit("content lost in conversion -> " + "; ".join(bad))
    print(f"  checked: {got_secs} sections, {got_figs} figures, {got_refs} references, {got_cards} annotation cards")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("source", nargs="?", default="thesis_source.html", type=Path)
    ap.add_argument("--no-annotations", action="store_true",
                    help="submission-only document: thesis text without the commentary")
    a = ap.parse_args()
    target = Path("Nischay_Maharjan_230357_Thesis_SubmissionOnly.docx" if a.no_annotations
                  else "Nischay_Maharjan_230357_Thesis_Annotated.docx")
    print(f"wrote {build(a.source, target, annotations=not a.no_annotations).resolve()}")
