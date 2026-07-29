"""Build the submission document in plain academic format.

The styled builder exists to make a document that reads well on screen: a sans
face for headings, an accent colour for section tags, varied sizes. Submission
usually wants the opposite, and this produces that instead.

  Times New Roman throughout, 12 pt, including headings and table text
  Body paragraphs justified
  Headings bold, and nothing else distinguishing them
  Automatic colour everywhere, so nothing is set to an explicit value

Content comes from the same annotated source as every other build, so the two
documents cannot drift apart. Annotations are excluded: this is the thesis.

CLI:  python -m scripts.build_thesis_plain_docx
"""
from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SRC = Path("thesis_source.html")
PNG = Path("reports/figures/thesis_png")
OUT = Path("Nischay_Maharjan_230357_Thesis.docx")

FONT = "Times New Roman"
SIZE = 12.0
HG = re.compile(r'<span class="hg[^"]*">(.*?)</span>', re.S)


# --------------------------------------------------------------- document set-up

def configure(doc: Document):
    """One font, one size, everywhere. Word inherits from Normal, so setting it
    here means nothing downstream has to remember to."""
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    # east-asian mapping too, or Word substitutes for some glyphs
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_after = Pt(10)

    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        s = doc.styles[name]
        s.font.name = FONT
        s.font.size = Pt(SIZE)
        s.font.bold = True
        s.font.italic = False
        s.font.color.rgb = None          # automatic
        s.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        # Word ships headings in a theme colour; clear it so they print black
        rpr = s.element.get_or_add_rPr()
        for tag in rpr.findall(qn("w:color")):
            rpr.remove(tag)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), "auto")
        rpr.append(col)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(8)
        s.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(1.0)
        sec.top_margin = sec.bottom_margin = Inches(1.0)


def page_numbers(doc: Document):
    """A PAGE field in the footer, centred."""
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    for el, attrs in (("w:fldChar", {"w:fldCharType": "begin"}),
                      ("w:instrText", None),
                      ("w:fldChar", {"w:fldCharType": "end"})):
        e = OxmlElement(el)
        if el == "w:instrText":
            e.set(qn("xml:space"), "preserve")
            e.text = " PAGE "
        for k, v in (attrs or {}).items():
            e.set(qn(k), v)
        r._r.append(e)
    r.font.name = FONT
    r.font.size = Pt(SIZE)


# ------------------------------------------------------------------- text runs

def text_of(s: str) -> str:
    s = re.sub(r"<br\s*/?>", " ", s)
    s = re.sub(r"<[^>]+>", "", s)
    return re.sub(r"\s+", " ", html.unescape(s)).strip()


def add_runs(par, frag: str):
    """Emit a fragment, keeping only bold and italic. Highlight spans are
    stripped: they carry meaning on screen and nothing on paper."""
    frag = HG.sub(r"\1", frag)
    frag = re.sub(r"<br\s*/?>", " ", frag)
    parts = re.split(r"(<(?:strong|b|em|i|code)[^>]*>.*?</(?:strong|b|em|i|code)>)",
                     frag, flags=re.S)
    for part in parts:
        if not part:
            continue
        m = re.match(r"<(strong|b|em|i|code)[^>]*>(.*?)</\1>", part, flags=re.S)
        if m:
            inner = text_of(m.group(2))
            if not inner:
                continue
            r = par.add_run(inner)
            r.bold = m.group(1) in ("strong", "b")
            r.italic = m.group(1) in ("em", "i")
        else:
            raw = html.unescape(re.sub(r"<[^>]+>", "", part))
            txt = re.sub(r"\s+", " ", raw)
            if not txt.strip():
                # a whitespace-only fragment between two inline elements is the
                # space that separates them; dropping it welds the words together
                if txt and par.runs:
                    par.add_run(" ")
                continue
            par.add_run(txt)


# ---------------------------------------------------------------------- blocks

def add_table(doc: Document, chunk: str):
    heads = re.findall(r"<th[^>]*>(.*?)</th>", chunk, flags=re.S)
    body = chunk[chunk.index("<tbody>"):] if "<tbody>" in chunk else chunk
    rows = [re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", r, flags=re.S)
            for r in re.findall(r"<tr[^>]*>(.*?)</tr>", body, flags=re.S)]
    rows = [r for r in rows if r]
    if not rows:
        return
    ncol = max(len(heads) or 0, max(len(r) for r in rows))
    tbl = doc.add_table(rows=0, cols=ncol)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if heads:
        cells = tbl.add_row().cells
        for i, h in enumerate(heads[:ncol]):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text_of(h))
            r.bold = True
    for row in rows:
        if heads and row == rows[0] and [text_of(c) for c in row] == [text_of(h) for h in heads]:
            continue
        cells = tbl.add_row().cells
        for i, c in enumerate(row[:ncol]):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, c)
    doc.add_paragraph()


def add_figure(doc: Document, caption: str):
    m = re.match(r"\s*(Figure|Table)\s*(\d+)", caption)
    if not m:
        return False
    png = PNG / f"{m.group(1)}{m.group(2)}.png"
    if not png.exists():
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(png), width=Inches(6.0))
    return True


def add_caption(doc: Document, text: str):
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    m = re.match(r"^((?:Figure|Table)\s*\d+)\.?\s*(.*)$", text, flags=re.S)
    if m:
        a = p.add_run(m.group(1) + ". ")
        a.bold = True
        p.add_run(m.group(2))
    else:
        p.add_run(text)


def emit_panel(doc: Document, panel: str, is_refs: bool):
    """Walk one thesis panel in document order."""
    blocks = re.split(r'(<figure>.*?</figure>|<div class="tablewrap">.*?</div>)',
                      panel, flags=re.S)
    for chunk in blocks:
        if not chunk or not chunk.strip():
            continue
        if chunk.startswith("<figure>"):
            cap = re.search(r"<figcaption>(.*?)</figcaption>", chunk, flags=re.S)
            cap_text = text_of(cap.group(1)) if cap else ""
            if add_figure(doc, cap_text):
                add_caption(doc, cap_text)
            continue
        if chunk.startswith('<div class="tablewrap">'):
            add_table(doc, chunk)
            continue
        for el in re.finditer(r"<(h4|p|ul|ol)[^>]*>(.*?)</\1>", chunk, flags=re.S):
            tag, inner = el.group(1), el.group(2)
            if tag == "h4":
                doc.add_paragraph(text_of(inner), style="Heading 2")
            elif tag == "p":
                if "tablewrap" in inner or not text_of(inner):
                    continue
                p = doc.add_paragraph()
                if is_refs:
                    # APA hanging indent
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(-0.5)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.line_spacing = 1.5
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_runs(p, inner)
            else:
                for li in re.findall(r"<li[^>]*>(.*?)</li>", inner, flags=re.S):
                    p = doc.add_paragraph(
                        style="List Bullet" if tag == "ul" else "List Number")
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    add_runs(p, li)


# ------------------------------------------------------------------------ build

def build() -> Path:
    raw = SRC.read_text(encoding="utf-8", errors="replace")
    doc = Document()
    configure(doc)
    page_numbers(doc)

    sections = re.findall(r'<section class="sec" id="([^"]+)">(.*?)</section>',
                          raw, flags=re.S)
    if not sections:
        raise SystemExit("no sections found in the source")

    for n, (sec_id, body) in enumerate(sections):
        title = re.search(r'class="sec-title">(.*?)</h2>', body, flags=re.S)
        if n:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        doc.add_paragraph(text_of(title.group(1)) if title else sec_id, style="Heading 1")

        is_refs = "thesis-panel refs" in body
        panel = re.search(r'<div class="thesis-panel[^"]*"[^>]*>(.*)', body, flags=re.S)
        if panel:
            content = panel.group(1)
            cut = content.find('<div class="ann-panel">')
            emit_panel(doc, content[:cut] if cut > 0 else content, is_refs)
        else:
            intro = re.search(r'<p style="[^"]*">(.*?)</p>', body, flags=re.S)
            if intro:
                add_runs(doc.add_paragraph(), intro.group(1))
            for fig in re.findall(r"<figure>.*?</figure>", body, flags=re.S):
                cap = re.search(r"<figcaption>(.*?)</figcaption>", fig, flags=re.S)
                cap_text = text_of(cap.group(1)) if cap else ""
                if add_figure(doc, cap_text):
                    add_caption(doc, cap_text)

    doc.save(OUT)
    verify(raw, OUT)
    return OUT


def verify(raw: str, out: Path):
    """Fail loudly if content did not arrive, or if formatting drifted."""
    d = Document(out)
    want_secs = len(re.findall(r'<section class="sec"', raw))
    got_secs = len([p for p in d.paragraphs if p.style.name == "Heading 1"])
    want_figs = len(re.findall(r"<figure>", raw))
    got_figs = len(d.inline_shapes)

    bad = []
    if got_secs < want_secs:
        bad.append(f"sections {want_secs} -> {got_secs}")
    if got_figs < want_figs:
        bad.append(f"figures {want_figs} -> {got_figs}")

    wrong_font = {r.font.name for p in d.paragraphs for r in p.runs
                  if r.font.name and r.font.name != FONT}
    wrong_size = {r.font.size.pt for p in d.paragraphs for r in p.runs
                  if r.font.size and r.font.size.pt != SIZE}
    coloured = [r.text[:20] for p in d.paragraphs for r in p.runs
                if r.font.color and r.font.color.rgb is not None]
    if wrong_font:
        bad.append(f"non-Times runs: {wrong_font}")
    if wrong_size:
        bad.append(f"non-12pt runs: {wrong_size}")
    if coloured:
        bad.append(f"{len(coloured)} coloured run(s)")
    if bad:
        raise SystemExit("plain-format check failed -> " + "; ".join(bad))

    body = [p for p in d.paragraphs
            if p.style.name == "Normal" and len(p.text.split()) > 25]
    just = sum(1 for p in body if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
               or p.alignment is None)
    print(f"  {got_secs} sections, {got_figs} figures, {len(d.tables)} tables")
    print(f"  {just}/{len(body)} body paragraphs justified, all runs "
          f"{FONT} {SIZE:.0f}pt, colour automatic")


if __name__ == "__main__":
    print(f"wrote {build().resolve()}")
